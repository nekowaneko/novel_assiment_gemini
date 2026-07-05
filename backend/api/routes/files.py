"""
檔案相關 API 端點
- 上傳 .docx 檔案並讀取內容
- 匯出文字為 .docx 檔案
"""

import os
import re
from io import BytesIO
from urllib.parse import quote
from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from docx import Document

router = APIRouter()


class ExportRequest(BaseModel):
    """匯出請求模型"""
    content: str = Field(..., description="要匯出的文字內容")
    filename: str = Field(default="my_novel.docx", description="檔案名稱")


class UploadResponse(BaseModel):
    """上傳回應模型"""
    content: str = Field(..., description="檔案文字內容")
    filename: str = Field(..., description="原始檔案名稱")
    character_count: int = Field(..., description="字元數")


@router.post("/upload", response_model=UploadResponse)
async def upload_docx(file: UploadFile = File(..., description="上傳的 .docx 檔案")):
    """
    上傳 .docx 檔案並讀取內容
    
    回傳檔案中的純文字內容。
    """
    # 檢查檔案格式
    if not file.filename or not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="僅支援 .docx 格式")
    
    try:
        # 讀取上傳的檔案
        content = await file.read()
        file_stream = BytesIO(content)
        
        # 使用 python-docx 讀取
        doc = Document(file_stream)
        full_text = [para.text for para in doc.paragraphs]
        text_content = "\n".join(full_text)
        
        return UploadResponse(
            content=text_content,
            filename=file.filename,
            character_count=len(text_content)
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"讀取檔案失敗：{str(e)}")


def strip_html_tags(html_content: str) -> str:
    """移除 HTML 標籤，返回純文字"""
    # 移除 HTML 標籤
    clean = re.sub(r'<[^>]+>', '', html_content)
    # 解碼 HTML 實體
    clean = clean.replace('&nbsp;', ' ')
    clean = clean.replace('&lt;', '<')
    clean = clean.replace('&gt;', '>')
    clean = clean.replace('&amp;', '&')
    return clean


@router.post("/export")
async def export_docx(request: ExportRequest):
    """
    匯出文字為 .docx 檔案
    
    將文字內容轉換為 .docx 格式並回傳下載。
    """
    if not request.content.strip():
        raise HTTPException(status_code=400, detail="請提供要匯出的內容")
    
    try:
        # 建立新文件
        doc = Document()
        
        # 清理 HTML 標籤並按段落分割
        clean_content = strip_html_tags(request.content)
        paragraphs = clean_content.split('\n')
        for para in paragraphs:
            if para.strip():  # 跳過空行
                doc.add_paragraph(para.strip())
        
        # 儲存到記憶體
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        
        # 對檔名進行 URL 編碼以支援中文
        encoded_filename = quote(request.filename, safe='')
        
        # 回傳檔案下載 (使用 RFC 5987 格式支援 UTF-8 檔名)
        return StreamingResponse(
            byte_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"匯出檔案失敗：{str(e)}")

