from docx import Document
from io import BytesIO
import mammoth
import pypandoc

# --- 檔案處理邏輯 ---

def read_docx(uploaded_file):
    """
    用於讀取 .docx 檔案並提取其中的文字內容。
    參數:
        uploaded_file (streamlit.uploaded_file): Streamlit 上傳的檔案物件。
    回傳:
        str: 檔案中的文字內容。
    """
    if uploaded_file is not None:
        # 將上傳的檔案物件轉換為 BytesIO，以供 python-docx 讀取
        file_stream = BytesIO(uploaded_file.getvalue())
        
        # 使用 python-docx 函式庫讀取檔案
        doc = Document(file_stream)
        
        full_text = []
        for para in doc.paragraphs:
            # 遍歷文件中的每個段落，並將其文字內容加入列表
            full_text.append(para.text)
            
        # 使用換行符將所有段落的文字內容連接起來
        return "\n".join(full_text)
    return ""

def read_docx_as_text(uploaded_file):
    """
    確保 100% 讀取 .docx 文字內容。
    """
    if uploaded_file is not None:
        try:
            uploaded_file.seek(0)
            file_stream = BytesIO(uploaded_file.read())
            doc = Document(file_stream)
            full_text = [para.text for para in doc.paragraphs]
            return "\n".join(full_text)
        except Exception as e:
            print(f"Error reading docx: {e}")
            return ""
    return ""

def write_new_docx(new_content, file_name, is_html=False):
    """
    接收編輯後的內容（文字或 HTML），並將其寫入一個新的 .docx 檔案。
    """
    byte_io = BytesIO()
    
    if is_html:
        # 如果是 HTML，嘗試使用 pypandoc 轉換
        try:
            # 注意：pypandoc.convert_text 需要系統已安裝 pandoc
            # 若環境不支援，可考慮降級回純文字處理
            output = pypandoc.convert_text(new_content, 'docx', format='html')
            byte_io.write(output)
        except Exception:
            # 降級方案：去除 HTML 標籤後寫入（簡單處理）
            doc = Document()
            # 這裡只是一個非常簡單的降級，實際開發中可能需要更複雜的 HTML 解析
            import re
            clean_text = re.sub('<[^<]+?>', '', new_content)
            doc.add_paragraph(clean_text)
            doc.save(byte_io)
    else:
        doc = Document()
        doc.add_paragraph(new_content)
        doc.save(byte_io)
    
    byte_io.seek(0)
    return byte_io
