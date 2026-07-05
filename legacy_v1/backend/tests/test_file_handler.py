import pytest
from io import BytesIO
from docx import Document
from src.utils.file_handler import read_docx_as_html, write_new_docx

class MockUploadedFile:
    def __init__(self, content, name="test.docx"):
        self.content = content
        self.name = name
    def getvalue(self):
        return self.content

def test_read_docx_as_html_basic():
    # Create a simple docx in memory
    doc = Document()
    doc.add_paragraph("Hello World")
    file_stream = BytesIO()
    doc.save(file_stream)
    
    mock_file = MockUploadedFile(file_stream.getvalue())
    html = read_docx_as_html(mock_file)
    
    assert "<p>Hello World</p>" in html

def test_read_docx_as_html_empty():
    assert read_docx_as_html(None) == ""

def test_write_new_docx_text():
    content = "New Content"
    result = write_new_docx(content, "test.docx", is_html=False)
    
    doc = Document(result)
    assert doc.paragraphs[0].text == "New Content"

def test_write_new_docx_html_fallback(mocker):
    # Mock pypandoc.convert_text to raise exception to test fallback
    mocker.patch("pypandoc.convert_text", side_effect=Exception("Pandoc not found"))
    
    html_content = "<p>HTML Content</p>"
    result = write_new_docx(html_content, "test.docx", is_html=True)
    
    doc = Document(result)
    # The fallback re.sub('<[^<]+?>', '', new_content)
    assert doc.paragraphs[0].text == "HTML Content"
