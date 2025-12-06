from docx import Document
from io import BytesIO

# --- 檔案處理邏輯 ---

def read_docx(uploaded_file):
    """
    用於讀取 .docx 檔案並提取其中的文字內容。
    參數:
        uploaded_file (streamlit.uploaded_file): Streamlit 上傳的檔案物件。
    回傳:
        str: 檔案中的文字內容。
    """
    if uploaded_file is not None:
        # 將上傳的檔案物件轉換為 BytesIO，以供 python-docx 讀取
        file_stream = BytesIO(uploaded_file.getvalue())
        
        # 使用 python-docx 函式庫讀取檔案
        doc = Document(file_stream)
        
        full_text = []
        for para in doc.paragraphs:
            # 遍歷文件中的每個段落，並將其文字內容加入列表
            full_text.append(para.text)
            
        # 使用換行符將所有段落的文字內容連接起來
        return "\n".join(full_text)
    return ""

def write_new_docx(new_content,file_name):
    """
    接收編輯後的文字內容，並將其寫入一個新的 .docx 檔案。
    參數:
        new_content (str): 編輯後的文字內容。
        file_name (str): 儲存新檔案的名稱。
    回傳:
        BytesIO: 包含新 .docx 檔案內容的位元流，以便在 Streamlit 中下載。
    """
    # 建立一個新的 Word 文件物件
    doc = Document()
    
    # 將所有文字內容添加到文件
    doc.add_paragraph(new_content)
    
    # 使用 BytesIO 建立一個在記憶體中的二進位流
    byte_io = BytesIO()
    doc.save(byte_io)
    
    # 將指針重置到文件的開頭，準備讀取
    byte_io.seek(0)
    
    return byte_io